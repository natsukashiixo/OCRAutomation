import docx
from docx.shared import Pt
from pathlib import Path
from bs4 import BeautifulSoup
from modules.functions_ui import ProgressCounter
from modules.logger_mod import write_log as WriteLog

output_file = Path('./TesseractOutput/to_be_parsed.docx')
hocr_xml = Path('./TesseractOutput/output.xml')

def parse_hocr(output_file=output_file, hocr_xml=hocr_xml):
    try:
        # open the hOCR file
        with open(hocr_xml, 'rb') as f:
            hocr_data = f.read()
    
        # parse the hOCR data using BeautifulSoup
        soup = BeautifulSoup(hocr_data, 'html.parser')
    
        # create a new Word document
        doc = docx.Document()
    
        # For new document (document-wide):
        # Set language value in the documents' default Run's Properties element.
        # I shamelessly stole this from StackOverflow
        styles_element = doc.styles.element
        rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
        lang_default = rpr_default.xpath('w:lang')[0]
        lang_default.set(docx.oxml.shared.qn('w:val'),'sv-SE')
        # This is also stolen from StackOverflow
        DOCX = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        element = doc.settings.element.find(DOCX + 'proofState')
        element.attrib[DOCX + 'grammar'] = 'dirty'
        element.attrib[DOCX + 'spelling'] = 'dirty'
    
        # Set the default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
    
        wordcount = list(soup.find_all(class_='ocrx_word'))
        print(f'Starting hOCR parsing. Found {len(wordcount)} words to process')
    
        with ProgressCounter(len(wordcount)) as progress:
            for page in soup.find_all(class_='ocr_page'):
                #for paragraph in page.find_all(class_='ocr_par'): #backup in case i break something
                for cArea in page.find_all(class_='ocr_carea'): #changed to use carea class instead of paragraph
                    p = doc.add_paragraph() 
                    for line in cArea.find_all(class_='ocr_line') + cArea.find_all(class_='ocr_caption'): 
                        linebreak = [] #if list > 1, does nothing
                        for word in line.find_all(class_='ocrx_word'): # iterate over the words in the line
                            # set the font and style of the text based on the formatting information
                            font = p.add_run().font
                            linebreak.append(word)
                            # check if the data-fonts attribute is present before accessing it
                            if 'data-fonts' in word.attrs:
                                font.name = word['data-fonts'].split()[0]
                            if 'data-font-size' in word.attrs:
                                font.size = docx.shared.Pt(float(word['data-font-size']))
                            if 'bold' in word['class']:
                                font.bold = True
                            if 'italic' in word['class']:
                                font.italic = True      
                            
                            # add the text to the paragraph
                            p.add_run(word.text)
                            p.add_run(' ')    
                            progress.update_progress()    
                        if len(linebreak) == 1:
                            p.add_run('\n')
                    
        doc.save(output_file)
        progress.finalize()
    except Exception as e:
        WriteLog(e)
    
    
if __name__ == "__main__":
    parse_hocr(output_file, hocr_xml)