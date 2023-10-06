import docx
from docx.shared import Pt
from modules.logger_mod import write_log as WriteLog

input_file = './TesseractOutput/less_mistakes.txt'
output_file = './TesseractOutput/output_FINAL.docx'

def rewrite_docx(input_file=input_file, output_file=output_file):
    try:
        print('Starting rewrite of docx file')
        doc = docx.Document()
    
        # For new document (document-wide):
        # Set language value in the documents' default Run's Properties element.
        # I shamelessly stole this from StackOverflow
        styles_element = doc.styles.element
        rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
        lang_default = rpr_default.xpath('w:lang')[0]
        lang_default.set(docx.oxml.shared.qn('w:val'),'sv-SE')
        # This is also stolen from StackOverflow
        DOCX = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        element = doc.settings.element.find(DOCX + 'proofState')
        element.attrib[DOCX + 'grammar'] = 'dirty'
        element.attrib[DOCX + 'spelling'] = 'dirty'
    
        # Set the default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
    
        with open(input_file, 'rt') as f:
            text = f.read()
            
        p = doc.add_paragraph()
        p.add_run(text)
        doc.save(output_file)
    
        print(f'Corrected document saved as {output_file} in the TesseractOutput folder.')
    except Exception as e:
        WriteLog(e)
    
    
if __name__ == "__main__":
    rewrite_docx(input_file, output_file)